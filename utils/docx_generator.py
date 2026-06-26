from docx import Document
import os

def write_inline_formatting_docx(paragraph, text):
    """Parses text split by ** to create alternating normal and bold runs."""
    parts = text.split('**')
    for idx, part in enumerate(parts):
        run = paragraph.add_run(part)
        if idx % 2 == 1:
            run.bold = True

def write_markdown_to_docx(doc, text):
    if not text:
        return
        
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Check if line is a header
        if line.startswith('#'):
            level = 0
            while level < len(line) and line[level] == '#':
                level += 1
            header_text = line[level:].strip()
            # Map level to docx heading level (e.g. level 1 to 3)
            doc.add_heading(header_text, level=min(level, 3))
            
        # Check if bullet point
        elif line.startswith('- ') or line.startswith('* '):
            bullet_text = line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            write_inline_formatting_docx(p, bullet_text)
            
        # Check if numbered list
        elif any(line.startswith(f"{i}. ") for i in range(1, 100)):
            dot_idx = line.find('. ')
            num = line[:dot_idx]
            item_text = line[dot_idx+2:].strip()
            p = doc.add_paragraph(style='List Number')
            write_inline_formatting_docx(p, item_text)
            
        else:
            p = doc.add_paragraph()
            write_inline_formatting_docx(p, line)

def create_docx_report(data: dict, output_path: str):
    doc = Document()
    
    doc.add_heading('Meeting Intelligence AI Report', 0)
    
    doc.add_paragraph(f"Customer: {data.get('customer_name', '')}")
    doc.add_paragraph(f"Date: {data.get('meeting_date', '')}")
    doc.add_paragraph(f"Sentiment: {data.get('sentiment', '')}")
    doc.add_paragraph(f"Deal Stage: {data.get('deal_stage', '')} (Probability: {data.get('probability', '')})")
    
    sections = [
        ("Minutes of Meeting (MoM)", data.get('mom')),
        ("Summary", data.get('summary')),
        ("Customer Insights", data.get('customer_insights')),
    ]
    
    for title, content in sections:
        if content:
            doc.add_heading(title, level=1)
            write_markdown_to_docx(doc, content)
            
    action_items = data.get('action_items', [])
    if action_items:
        doc.add_heading('Action Items', level=1)
        for ai in action_items:
            if isinstance(ai, dict):
                p = doc.add_paragraph(style='List Bullet')
                text = f"**{ai.get('task')}** (Owner: {ai.get('owner')}, Due: {ai.get('due_date')})"
                write_inline_formatting_docx(p, text)
            else:
                p = doc.add_paragraph(style='List Bullet')
                write_inline_formatting_docx(p, str(ai))
            
    doc.save(output_path)
    return output_path
